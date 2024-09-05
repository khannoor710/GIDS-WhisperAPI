from flask import Flask, request, send_file, jsonify
from pdf2docx import Converter
from docx import Document
from transformers import MarianMTModel, MarianTokenizer
import torch
import os
import logging
from docx2pdf import convert

app = Flask(__name__)

# Set up logging
logging.basicConfig(level=logging.DEBUG)

# Load MarianMT model and tokenizer for French to English translation
model_name = 'Helsinki-NLP/opus-mt-fr-en'
tokenizer = MarianTokenizer.from_pretrained(model_name)
model = MarianMTModel.from_pretrained(model_name)

device = "cuda" if torch.cuda.is_available() else "cpu"
model = model.to(device)

def translate_text(text):
    """Translate French text to English using MarianMT model."""
    batch = tokenizer([text], return_tensors="pt", padding=True, truncation=True, max_length=512).to(device)
    translated = model.generate(**batch)
    return tokenizer.decode(translated[0], skip_special_tokens=True)

def convert_pdf_to_word(pdf_path, word_path):
    """Convert PDF file to Word format."""
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

def translate_docx(docx_path):
    """Translate all text in a Word document from French to English."""
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            translated_text = translate_text(paragraph.text)
            paragraph.text = translated_text

    # Translate text in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    translated_text = translate_text(cell.text)
                    cell.text = translated_text

    doc.save(docx_path)

def convert_word_to_pdf(docx_path, pdf_path):
    """Convert Word document to PDF format."""
    convert(docx_path, pdf_path)

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload, translation, and conversion back to PDF."""
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and file.filename.lower().endswith('.pdf'):
        input_pdf_path = os.path.join(os.getcwd(), 'input.pdf')
        intermediate_docx_path = os.path.join(os.getcwd(), 'intermediate.docx')
        output_pdf_path = os.path.join(os.getcwd(), 'translated_english_pdf.pdf')

        try:
            file.save(input_pdf_path)

            # Step 1: Convert PDF to Word
            convert_pdf_to_word(input_pdf_path, intermediate_docx_path)

            # Step 2: Translate text in Word document
            translate_docx(intermediate_docx_path)

            # Step 3: Convert translated Word document back to PDF
            convert_word_to_pdf(intermediate_docx_path, output_pdf_path)

            # Clean up intermediate files
            os.remove(input_pdf_path)
            os.remove(intermediate_docx_path)

            if os.path.exists(output_pdf_path):
                response = send_file(output_pdf_path, as_attachment=True)
                os.remove(output_pdf_path)  # Clean up after sending
                return response
            else:
                return jsonify({"error": "Translation failed, file not found."}), 500

        except Exception as e:
            logging.error(f"Error in upload_file: {e}")
            return jsonify({"error": "Error during translation."}), 500

    return jsonify({"error": "Invalid file format"}), 400

if __name__ == '__main__':
    app.run(debug=True)